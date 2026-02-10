#!/usr/bin/env python3
"""
Template Manager Script for Report Generator

Manages the template library including listing, validating,
recommending, and managing presets.

Usage:
    python template_manager.py --list
    python template_manager.py --info template.docx
    python template_manager.py --recommend "地质调查报告"
    python template_manager.py --save-preset "my-format"
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class TemplateManager:
    """Manages the template library."""
    
    def __init__(self, templates_dir: str = "templates"):
        self.templates_dir = Path(templates_dir)
        self.presets_dir = self.templates_dir / ".presets"
        self.presets_dir.mkdir(parents=True, exist_ok=True)
    
    def list_templates(self) -> List[Dict]:
        """List all available templates."""
        templates = []
        
        if not self.templates_dir.exists():
            print(f"模板目录不存在: {self.templates_dir}")
            return templates
        
        for template_file in self.templates_dir.glob("*.docx"):
            template_info = self._get_template_info(template_file)
            templates.append(template_info)
        
        return templates
    
    def _get_template_info(self, template_path: Path) -> Dict:
        """Get information about a template."""
        info = {
            'name': template_path.stem,
            'path': str(template_path),
            'size': template_path.stat().st_size,
            'modified': datetime.fromtimestamp(template_path.stat().st_mtime).strftime('%Y-%m-%d %H:%M')
        }
        
        # Try to read document info
        if DOCX_AVAILABLE:
            try:
                doc = Document(template_path)
                info['paragraphs'] = len(doc.paragraphs)
                info['tables'] = len(doc.tables)
                
                # Get styles
                styles = [s.name for s in doc.styles if s.type == 1]  # Paragraph styles
                info['styles'] = styles[:10]  # First 10 styles
            except Exception as e:
                info['error'] = str(e)
        
        return info
    
    def show_template_info(self, template_name: str):
        """Show detailed information about a template."""
        template_path = self.templates_dir / template_name
        
        if not template_path.exists():
            print(f"错误: 模板不存在: {template_name}")
            return
        
        info = self._get_template_info(template_path)
        
        print("\n" + "="*60)
        print(f"模板信息: {info['name']}")
        print("="*60)
        print(f"路径: {info['path']}")
        print(f"大小: {info['size'] / 1024:.1f} KB")
        print(f"修改时间: {info['modified']}")
        
        if 'paragraphs' in info:
            print(f"段落数: {info['paragraphs']}")
        if 'tables' in info:
            print(f"表格数: {info['tables']}")
        if 'styles' in info:
            print(f"样式: {', '.join(info['styles'])}...")
        
        if 'error' in info:
            print(f"警告: 读取模板时出错 - {info['error']}")
        
        print("="*60)
    
    def recommend_template(self, report_type: str) -> Optional[str]:
        """Recommend a template based on report type."""
        # Define recommendation rules
        recommendations = {
            '地质': 'geological-survey-template.docx',
            '地质调查': 'geological-survey-template.docx',
            '矿产': 'geological-survey-template.docx',
            '工程': 'generic-engineering-template.docx',
            '技术': 'generic-engineering-template.docx',
            '设计': 'generic-engineering-template.docx',
            '项目': 'generic-engineering-template.docx'
        }
        
        # Find matching recommendation
        report_type_lower = report_type.lower()
        for keyword, template in recommendations.items():
            if keyword in report_type_lower:
                template_path = self.templates_dir / template
                if template_path.exists():
                    return str(template_path)
        
        # Return first available template
        templates = self.list_templates()
        if templates:
            return templates[0]['path']
        
        return None
    
    def validate_templates(self) -> List[Dict]:
        """Validate all templates in the library."""
        results = []
        
        templates = self.list_templates()
        for template in templates:
            result = {'name': template['name'], 'valid': True, 'issues': []}
            
            if 'error' in template:
                result['valid'] = False
                result['issues'].append(f"无法读取: {template['error']}")
            
            if template['size'] < 1000:
                result['issues'].append("文件过小，可能不完整")
            
            results.append(result)
        
        return results
    
    def delete_template(self, template_name: str) -> bool:
        """Delete a template."""
        template_path = self.templates_dir / template_name
        
        if not template_path.exists():
            print(f"错误: 模板不存在: {template_name}")
            return False
        
        try:
            template_path.unlink()
            print(f"✓ 已删除模板: {template_name}")
            return True
        except Exception as e:
            print(f"错误: 删除失败 - {e}")
            return False
    
    def save_preset(self, name: str, config: Dict):
        """Save a configuration as a named preset."""
        preset_file = self.presets_dir / f"{name}.json"
        
        try:
            with open(preset_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=2)
            print(f"✓ 预设已保存: {name}")
            return True
        except Exception as e:
            print(f"错误: 保存预设失败 - {e}")
            return False
    
    def load_preset(self, name: str) -> Optional[Dict]:
        """Load a named preset."""
        preset_file = self.presets_dir / f"{name}.json"
        
        if not preset_file.exists():
            return None
        
        try:
            with open(preset_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"错误: 加载预设失败 - {e}")
            return None
    
    def list_presets(self) -> List[str]:
        """List all available presets."""
        if not self.presets_dir.exists():
            return []
        
        presets = []
        for preset_file in self.presets_dir.glob("*.json"):
            presets.append(preset_file.stem)
        
        return presets
    
    def show_preset_info(self, name: str):
        """Show information about a preset."""
        config = self.load_preset(name)
        
        if config is None:
            print(f"错误: 预设不存在: {name}")
            return
        
        print("\n" + "="*60)
        print(f"预设信息: {name}")
        print("="*60)
        
        if 'page' in config:
            page = config['page']
            print(f"页边距: 上{page.get('margin_top', 'N/A')}/下{page.get('margin_bottom', 'N/A')}/"
                  f"左{page.get('margin_left', 'N/A')}/右{page.get('margin_right', 'N/A')} cm")
            print(f"装订线: {page.get('gutter', 'N/A')} cm")
        
        if 'fonts' in config:
            fonts = config['fonts']
            print(f"中文字体: {fonts.get('chinese', 'N/A')}")
            print(f"章标题字体: {fonts.get('chapter_title', 'N/A')}")
        
        if 'sizes' in config:
            sizes = config['sizes']
            print(f"章标题字号: {sizes.get('chapter', 'N/A')} pt")
            print(f"正文字号: {sizes.get('body', 'N/A')} pt")
        
        print("="*60)


def print_template_list(templates: List[Dict]):
    """Print a formatted list of templates."""
    if not templates:
        print("\n没有找到模板文件")
        return
    
    print("\n" + "="*80)
    print(f"{'序号':<6}{'模板名称':<30}{'大小':<12}{'修改时间':<20}")
    print("="*80)
    
    for i, template in enumerate(templates, 1):
        size_kb = template['size'] / 1024
        print(f"{i:<6}{template['name']:<30}{size_kb:>6.1f} KB{template['modified']:>20}")
    
    print("="*80)
    print(f"总计: {len(templates)} 个模板\n")


def print_validation_results(results: List[Dict]):
    """Print validation results."""
    print("\n" + "="*60)
    print("模板验证结果")
    print("="*60)
    
    for result in results:
        status = "✓ 有效" if result['valid'] else "✗ 无效"
        print(f"{result['name']}: {status}")
        
        if result['issues']:
            for issue in result['issues']:
                print(f"  - {issue}")
    
    print("="*60)


def main():
    parser = argparse.ArgumentParser(
        description="管理Word模板库",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 列出所有模板
  python template_manager.py --list
  
  # 查看模板详情
  python template_manager.py --info geological-survey-template.docx
  
  # 根据报告类型推荐模板
  python template_manager.py --recommend "地质调查报告"
  
  # 验证所有模板
  python template_manager.py --validate
  
  # 删除模板
  python template_manager.py --delete old-template.docx
  
  # 列出所有预设
  python template_manager.py --list-presets
  
  # 查看预设详情
  python template_manager.py --preset-info "my-custom-format"
        """
    )
    
    parser.add_argument('--list', '-l', action='store_true',
                       help='列出所有模板')
    parser.add_argument('--info', metavar='TEMPLATE',
                       help='显示模板详细信息')
    parser.add_argument('--recommend', metavar='TYPE',
                       help='根据报告类型推荐模板')
    parser.add_argument('--validate', action='store_true',
                       help='验证所有模板')
    parser.add_argument('--delete', metavar='TEMPLATE',
                       help='删除指定模板')
    parser.add_argument('--list-presets', action='store_true',
                       help='列出所有预设')
    parser.add_argument('--preset-info', metavar='NAME',
                       help='显示预设详细信息')
    parser.add_argument('--templates-dir', default='templates',
                       help='模板目录路径 (默认: templates)')
    
    args = parser.parse_args()
    
    # Create manager
    manager = TemplateManager(args.templates_dir)
    
    # Execute command
    if args.list:
        templates = manager.list_templates()
        print_template_list(templates)
    
    elif args.info:
        manager.show_template_info(args.info)
    
    elif args.recommend:
        template = manager.recommend_template(args.recommend)
        if template:
            print(f"\n推荐模板: {template}")
            print(f"报告类型: {args.recommend}")
        else:
            print(f"\n未找到适合 '{args.recommend}' 的模板")
            print("建议生成新模板或使用通用模板")
    
    elif args.validate:
        results = manager.validate_templates()
        print_validation_results(results)
    
    elif args.delete:
        manager.delete_template(args.delete)
    
    elif args.list_presets:
        presets = manager.list_presets()
        if presets:
            print("\n可用预设:")
            for i, preset in enumerate(presets, 1):
                print(f"  {i}. {preset}")
        else:
            print("\n没有保存的预设")
    
    elif args.preset_info:
        manager.show_preset_info(args.preset_info)
    
    else:
        parser.print_help()
    
    return 0


if __name__ == '__main__':
    sys.exit(main())
