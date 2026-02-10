#!/usr/bin/env python3
"""
Apply Template Script for Report Generator

Applies Word document templates to existing documents.
Supports three formatting methods: mapping, copy, and partial.

Usage:
    python apply_template.py --input source.docx --template template.docx --output result.docx --method mapping
    python apply_template.py --input source.docx --template template.docx --output result.docx --method copy
    python apply_template.py --input source.docx --template template.docx --output result.docx --method partial --parts pagesetup,body,table
"""

import argparse
import os
import sys
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from copy import deepcopy

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Error: python-docx library not found. Install with: pip install python-docx")
    sys.exit(1)


class StyleMapper:
    """Maps document styles based on paragraph characteristics."""
    
    def __init__(self, template_doc: Document):
        self.template_doc = template_doc
        self.style_mapping = {}
        self._analyze_template_styles()
    
    def _analyze_template_styles(self):
        """Analyze template styles to create mapping rules."""
        # Define style characteristics based on template
        self.style_rules = {
            'ChapterTitle': {
                'font_size_range': (20, 26),  # 22pt ± 4pt
                'bold': True,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,
                'priority': 1
            },
            'SectionTitle': {
                'font_size_range': (14, 18),  # 16pt ± 2pt
                'bold': True,
                'alignment': WD_ALIGN_PARAGRAPH.LEFT,
                'priority': 2
            },
            'SubsectionTitle': {
                'font_size_range': (12, 16),  # 14pt ± 2pt
                'bold': True,
                'alignment': WD_ALIGN_PARAGRAPH.LEFT,
                'priority': 3
            },
            'BodyText': {
                'font_size_range': (9, 12),  # 10.5pt ± 1.5pt
                'bold': False,
                'alignment': WD_ALIGN_PARAGRAPH.JUSTIFY,
                'priority': 4
            }
        }
    
    def identify_paragraph_style(self, paragraph) -> str:
        """Identify the style of a paragraph based on its characteristics."""
        if not paragraph.runs:
            return 'BodyText'
        
        # Get paragraph characteristics
        first_run = paragraph.runs[0]
        font_size = first_run.font.size.pt if first_run.font.size else 10.5
        is_bold = first_run.font.bold
        alignment = paragraph.alignment
        
        # Match against rules
        scores = {}
        for style_name, rules in self.style_rules.items():
            score = 0
            
            # Check font size
            if rules['font_size_range'][0] <= font_size <= rules['font_size_range'][1]:
                score += 3
            
            # Check bold
            if rules['bold'] == is_bold:
                score += 2
            
            # Check alignment
            if rules['alignment'] == alignment:
                score += 1
            
            scores[style_name] = score
        
        # Return style with highest score
        if scores:
            best_match = max(scores, key=scores.get)
            if scores[best_match] > 0:
                return best_match
        
        return 'BodyText'


class TemplateApplier:
    """Applies template formatting to existing documents."""
    
    def __init__(self, source_path: str, template_path: str, output_path: str):
        self.source_path = Path(source_path)
        self.template_path = Path(template_path)
        self.output_path = Path(output_path)
        self.source_doc = None
        self.template_doc = None
        self.output_doc = None
        
    def load_documents(self):
        """Load source and template documents."""
        print(f"加载源文档: {self.source_path}")
        self.source_doc = Document(self.source_path)
        
        print(f"加载模板: {self.template_path}")
        self.template_doc = Document(self.template_path)
        
    def apply_mapping_method(self):
        """Apply template using style mapping method (recommended)."""
        print("\n使用样式映射方法...")
        
        # Create new document from template
        self.output_doc = deepcopy(self.template_doc)
        
        # Clear sample content but keep styles
        # We'll rebuild the document with source content
        
        # Create style mapper
        mapper = StyleMapper(self.template_doc)
        
        # Statistics
        stats = {
            'ChapterTitle': 0,
            'SectionTitle': 0,
            'SubsectionTitle': 0,
            'BodyText': 0,
            'tables': 0,
            'figures': 0
        }
        
        # Process each paragraph in source
        for para in self.source_doc.paragraphs:
            # Identify style
            style_name = mapper.identify_paragraph_style(para)
            stats[style_name] += 1
            
            # Create new paragraph in output with identified style
            new_para = self.output_doc.add_paragraph(para.text, style=style_name)
            
            # Copy formatting if needed
            self._copy_paragraph_format(para, new_para)
        
        # Process tables
        for table in self.source_doc.tables:
            stats['tables'] += 1
            self._copy_table(table)
        
        return stats
    
    def apply_copy_method(self):
        """Apply template using full copy method."""
        print("\n使用全文复制方法...")
        
        # Start with template
        self.output_doc = deepcopy(self.template_doc)
        
        # Clear all paragraphs
        for para in list(self.output_doc.paragraphs):
            p = para._element
            p.getparent().remove(p)
        
        # Copy content from source with template styles
        mapper = StyleMapper(self.template_doc)
        
        stats = {'paragraphs': 0, 'tables': 0}
        
        for para in self.source_doc.paragraphs:
            style_name = mapper.identify_paragraph_style(para)
            new_para = self.output_doc.add_paragraph(para.text, style=style_name)
            stats['paragraphs'] += 1
        
        for table in self.source_doc.tables:
            self._copy_table(table)
            stats['tables'] += 1
        
        return stats
    
    def apply_partial_method(self, parts: List[str]):
        """Apply only specific parts of template."""
        print(f"\n使用局部应用方法: {', '.join(parts)}")
        
        # Start with source document
        self.output_doc = deepcopy(self.source_doc)
        
        stats = {}
        
        if 'pagesetup' in parts:
            print("  应用页面设置...")
            self._apply_page_setup()
            stats['page_setup'] = 'applied'
        
        if 'body' in parts:
            print("  应用正文样式...")
            body_count = self._apply_body_style()
            stats['body_paragraphs'] = body_count
        
        if 'table' in parts:
            print("  应用表格样式...")
            table_count = self._apply_table_styles()
            stats['tables'] = table_count
        
        if 'figure' in parts:
            print("  应用图片样式...")
            figure_count = self._apply_figure_styles()
            stats['figures'] = figure_count
        
        return stats
    
    def _copy_paragraph_format(self, source_para, target_para):
        """Copy paragraph formatting."""
        # Copy alignment
        if source_para.alignment:
            target_para.alignment = source_para.alignment
        
        # Copy other formatting as needed
        # Note: Most formatting comes from styles
    
    def _copy_table(self, source_table):
        """Copy a table from source to output."""
        # Create table with same dimensions
        rows = len(source_table.rows)
        cols = len(source_table.columns)
        
        new_table = self.output_doc.add_table(rows=rows, cols=cols)
        new_table.style = 'Table Grid'
        
        # Copy content
        for i, row in enumerate(source_table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text
    
    def _apply_page_setup(self):
        """Apply page setup from template."""
        template_section = self.template_doc.sections[0]
        output_section = self.output_doc.sections[0]
        
        output_section.page_height = template_section.page_height
        output_section.page_width = template_section.page_width
        output_section.left_margin = template_section.left_margin
        output_section.right_margin = template_section.right_margin
        output_section.top_margin = template_section.top_margin
        output_section.bottom_margin = template_section.bottom_margin
        output_section.gutter = template_section.gutter
        output_section.header_distance = template_section.header_distance
        output_section.footer_distance = template_section.footer_distance
    
    def _apply_body_style(self) -> int:
        """Apply body text style to all body paragraphs."""
        count = 0
        for para in self.output_doc.paragraphs:
            # Check if it's a body paragraph (not heading)
            if not para.style.name.startswith('Heading'):
                try:
                    para.style = self.output_doc.styles['BodyText']
                    count += 1
                except:
                    pass
        return count
    
    def _apply_table_styles(self) -> int:
        """Apply table styles."""
        count = 0
        for table in self.output_doc.tables:
            table.style = 'Table Grid'
            count += 1
        return count
    
    def _apply_figure_styles(self) -> int:
        """Apply figure styles."""
        count = 0
        # Implementation depends on how figures are identified
        return count
    
    def save(self):
        """Save the output document."""
        # Ensure output directory exists
        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        
        print(f"\n保存结果: {self.output_path}")
        self.output_doc.save(self.output_path)
        print("✓ 完成")


def print_statistics(stats: Dict, method: str):
    """Print formatting statistics."""
    print("\n" + "="*60)
    print("排版统计")
    print("="*60)
    
    if method == 'mapping':
        print(f"章标题: {stats.get('ChapterTitle', 0)} 个")
        print(f"节标题: {stats.get('SectionTitle', 0)} 个")
        print(f"段标题: {stats.get('SubsectionTitle', 0)} 个")
        print(f"正文段落: {stats.get('BodyText', 0)} 段")
        print(f"表格: {stats.get('tables', 0)} 个")
    elif method == 'copy':
        print(f"段落: {stats.get('paragraphs', 0)} 个")
        print(f"表格: {stats.get('tables', 0)} 个")
    elif method == 'partial':
        for key, value in stats.items():
            print(f"{key}: {value}")
    
    print("="*60)


def main():
    parser = argparse.ArgumentParser(
        description="将模板应用到已有Word文档",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 样式映射方法（推荐）- 自动识别标题层级
  python apply_template.py --input report.docx --template geo-template.docx --output formatted.docx --method mapping
  
  # 全文复制方法 - 完全套用模板
  python apply_template.py --input report.docx --template geo-template.docx --output formatted.docx --method copy
  
  # 局部应用方法 - 仅应用特定部分
  python apply_template.py --input report.docx --template geo-template.docx --output formatted.docx --method partial --parts pagesetup,body,table
        """
    )
    
    parser.add_argument('--input', '-i', required=True,
                       help='源文档路径')
    parser.add_argument('--template', '-t', required=True,
                       help='模板文件路径')
    parser.add_argument('--output', '-o', required=True,
                       help='输出文件路径')
    parser.add_argument('--method', '-m', 
                       choices=['mapping', 'copy', 'partial'],
                       default='mapping',
                       help='排版方法 (默认: mapping)')
    parser.add_argument('--parts',
                       help='局部应用时选择的部分，逗号分隔: pagesetup,body,table,figure')
    
    args = parser.parse_args()
    
    # Validate input files
    if not Path(args.input).exists():
        print(f"错误: 源文档不存在: {args.input}")
        return 1
    
    if not Path(args.template).exists():
        print(f"错误: 模板文件不存在: {args.template}")
        return 1
    
    # Parse parts for partial method
    parts = []
    if args.method == 'partial' and args.parts:
        parts = [p.strip() for p in args.parts.split(',')]
    
    # Create applier
    try:
        applier = TemplateApplier(args.input, args.template, args.output)
        applier.load_documents()
        
        # Apply template
        if args.method == 'mapping':
            stats = applier.apply_mapping_method()
        elif args.method == 'copy':
            stats = applier.apply_copy_method()
        elif args.method == 'partial':
            if not parts:
                print("错误: 局部应用方法需要指定 --parts 参数")
                return 1
            stats = applier.apply_partial_method(parts)
        
        # Save result
        applier.save()
        
        # Print statistics
        print_statistics(stats, args.method)
        
        print(f"\n✓ 排版完成: {args.output}")
        
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()
        return 1
    
    return 0


if __name__ == '__main__':
    sys.exit(main())
