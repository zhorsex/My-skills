#!/usr/bin/env python3
"""
Material Image Extraction Script

从 docx 格式素材文件中提取所有配图，建立配图与段落的关联关系，
生成素材配图索引文件。

Usage:
    python3 extract_docx_images.py --help
    python3 extract_docx_images.py --material F:/docs/素材.docx --output assets/docx_images/
    python3 extract_docx_images.py --material-dir F:/docs/ --output assets/docx_images/

Examples:
    extract_docx_images.py --material 生态修复技术导则.docx
    extract_docx_images.py --material-dir ./materials/ --verbose

Supports:
    - Word文档 (.docx) 图片提取
    - 图片格式: PNG, JPG, JPEG, EMF, WMF, BMP, GIF
    - 自动识别图片-段落关联
    - 生成素材配图索引 (material-index.md)
"""

import os
import sys
import argparse
import zipfile
import json
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple, Optional
import re

# Try to import docx library
try:
    from docx import Document
    from docx.oxml import parse_xml
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Using fallback method.")

# Try to import PIL for image processing
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("Warning: PIL/Pillow not available. Image processing features disabled.")


class MaterialImageExtractor:
    """素材配图提取器"""
    
    def __init__(self, output_dir: str, verbose: bool = False):
        self.output_dir = Path(output_dir)
        self.verbose = verbose
        self.extracted_images = []
        self.material_index = {}
        
        # 创建输出目录
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # 支持的图片格式
        self.supported_formats = {
            '.png': 'PNG',
            '.jpg': 'JPEG',
            '.jpeg': 'JPEG',
            '.emf': 'EMF',
            '.wmf': 'WMF',
            '.bmp': 'BMP',
            '.gif': 'GIF'
        }
    
    def log(self, message: str, level: str = "INFO"):
        """打印日志信息"""
        if self.verbose or level in ["ERROR", "WARNING"]:
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            print(f"[{timestamp}] [{level}] {message}")
    
    def extract_images_from_docx(self, docx_path: str) -> List[Dict]:
        """
        从 docx 文件中提取所有图片
        
        Args:
            docx_path: docx 文件路径
            
        Returns:
            提取的图片信息列表
        """
        self.log(f"开始提取: {docx_path}")
        
        images = []
        docx_path = Path(docx_path)
        
        if not docx_path.exists():
            self.log(f"文件不存在: {docx_path}", "ERROR")
            return images
        
        try:
            # 方法1: 使用 python-docx (如果有)
            if DOCX_AVAILABLE:
                images = self._extract_with_docx(docx_path)
            
            # 方法2: 使用 zipfile 解压 (更可靠，作为备选)
            if not images:
                images = self._extract_with_zipfile(docx_path)
                
        except Exception as e:
            self.log(f"提取失败: {str(e)}", "ERROR")
        
        self.log(f"提取完成: {len(images)} 张图片")
        return images
    
    def _extract_with_docx(self, docx_path: Path) -> List[Dict]:
        """使用 python-docx 提取图片"""
        images = []
        
        try:
            doc = Document(str(docx_path))
            
            # 遍历所有关系，查找图片
            rels = doc.part.rels
            for rel in rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_part = doc.part.related_parts[rel.target_ref]
                        image_data = image_part.blob
                        
                        # 确定图片格式
                        content_type = image_part.content_type
                        ext = self._get_extension_from_content_type(content_type)
                        
                        image_info = {
                            'data': image_data,
                            'format': ext,
                            'content_type': content_type,
                            'method': 'docx'
                        }
                        images.append(image_info)
                        
                    except Exception as e:
                        self.log(f"提取单张图片失败: {str(e)}", "WARNING")
                        
        except Exception as e:
            self.log(f"docx 方法失败: {str(e)}", "WARNING")
        
        return images
    
    def _extract_with_zipfile(self, docx_path: Path) -> List[Dict]:
        """使用 zipfile 解压提取图片 (更可靠的方法)"""
        images = []
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as z:
                # docx 文件结构: word/media/ 目录包含所有图片
                image_files = [f for f in z.namelist() if f.startswith('word/media/')]
                
                for img_file in image_files:
                    try:
                        image_data = z.read(img_file)
                        
                        # 从文件名获取扩展名
                        ext = Path(img_file).suffix.lower()
                        
                        image_info = {
                            'data': image_data,
                            'format': ext,
                            'original_path': img_file,
                            'method': 'zipfile'
                        }
                        images.append(image_info)
                        
                    except Exception as e:
                        self.log(f"读取图片失败 {img_file}: {str(e)}", "WARNING")
                        
        except Exception as e:
            self.log(f"zipfile 方法失败: {str(e)}", "ERROR")
        
        return images
    
    def _get_extension_from_content_type(self, content_type: str) -> str:
        """根据 content type 获取文件扩展名"""
        type_map = {
            'image/png': '.png',
            'image/jpeg': '.jpg',
            'image/jpg': '.jpg',
            'image/bmp': '.bmp',
            'image/gif': '.gif',
            'image/x-emf': '.emf',
            'image/x-wmf': '.wmf'
        }
        return type_map.get(content_type, '.png')
    
    def identify_paragraph_relationship(self, docx_path: str) -> Dict:
        """
        识别图片与段落的关联关系
        
        Args:
            docx_path: docx 文件路径
            
        Returns:
            图片-段落关联字典
        """
        relationships = {}
        
        if not DOCX_AVAILABLE:
            self.log("无法识别段落关联: python-docx 不可用", "WARNING")
            return relationships
        
        try:
            doc = Document(str(docx_path))
            
            # 遍历所有段落，记录包含图片的段落
            for para_idx, para in enumerate(doc.paragraphs):
                para_text = para.text.strip()
                
                # 检查段落是否包含图片
                # 注意：这里使用简化方法，实际可能需要更复杂的 XML 解析
                if para._element.xpath('.//a:blip', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'}):
                    relationships[f"para_{para_idx}"] = {
                        'text': para_text[:100],  # 前100字符
                        'index': para_idx,
                        'has_image': True
                    }
                    
        except Exception as e:
            self.log(f"段落关联识别失败: {str(e)}", "WARNING")
        
        return relationships
    
    def classify_image(self, image_data: bytes, filename: str) -> str:
        """
        对图片进行自动分类
        
        Args:
            image_data: 图片二进制数据
            filename: 原始文件名
            
        Returns:
            图片类型: 示意图, 流程图, 数据图, 照片, 其他
        """
        # 基于文件名关键词分类
        filename_lower = filename.lower()
        
        # 流程图关键词
        flowchart_keywords = ['流程', 'flow', 'process', '步骤', 'step', 'procedure']
        # 示意图关键词
        diagram_keywords = ['示意', '原理', '结构', '架构', 'diagram', 'schematic']
        # 数据图关键词
        data_keywords = ['数据', '统计', '图表', '对比', 'chart', 'data', 'graph']
        # 照片关键词
        photo_keywords = ['照片', '实景', '现场', 'photo', 'image', 'picture']
        
        for keyword in flowchart_keywords:
            if keyword in filename_lower:
                return '流程图'
        
        for keyword in data_keywords:
            if keyword in filename_lower:
                return '数据图'
        
        for keyword in diagram_keywords:
            if keyword in filename_lower:
                return '示意图'
        
        for keyword in photo_keywords:
            if keyword in filename_lower:
                return '照片'
        
        # 如果无法通过文件名判断，尝试通过内容判断
        if PIL_AVAILABLE:
            try:
                img = Image.open(io.BytesIO(image_data))
                width, height = img.size
                
                # 基于尺寸和模式的简单分类
                if img.mode == 'RGB' and width > 1000:
                    return '照片'
                elif width < 500 and height < 500:
                    return '图标'
                    
            except Exception:
                pass
        
        return '其他'
    
    def save_images(self, images: List[Dict], material_name: str) -> List[Dict]:
        """
        保存提取的图片到输出目录
        
        Args:
            images: 提取的图片列表
            material_name: 素材文件名（用于命名）
            
        Returns:
            保存的图片信息列表
        """
        saved_images = []
        material_dir = self.output_dir / material_name
        material_dir.mkdir(parents=True, exist_ok=True)
        
        for idx, img_info in enumerate(images, 1):
            try:
                # 确定文件扩展名
                ext = img_info.get('format', '.png')
                if not ext.startswith('.'):
                    ext = '.' + ext
                
                # 转换为标准格式（如果不是 PNG/JPG）
                if ext not in ['.png', '.jpg', '.jpeg'] and PIL_AVAILABLE:
                    img_info = self._convert_image_format(img_info)
                    ext = img_info['format']
                
                # 生成文件名
                filename = f"asset_{material_name}_{idx:03d}{ext}"
                filepath = material_dir / filename
                
                # 保存图片
                with open(filepath, 'wb') as f:
                    f.write(img_info['data'])
                
                # 获取图片信息
                img_metadata = self._get_image_metadata(filepath, img_info['data'])
                
                saved_info = {
                    'id': f"IMG_{material_name}_{idx:03d}",
                    'filename': filename,
                    'filepath': str(filepath.relative_to(self.output_dir.parent)),
                    'original_format': img_info.get('format', 'unknown'),
                    'size_bytes': len(img_info['data']),
                    'type': self.classify_image(img_info['data'], filename),
                    **img_metadata
                }
                saved_images.append(saved_info)
                
                self.log(f"保存图片: {filename} ({saved_info['type']})")
                
            except Exception as e:
                self.log(f"保存图片失败 #{idx}: {str(e)}", "ERROR")
        
        return saved_images
    
    def _convert_image_format(self, img_info: Dict, target_format: str = '.png') -> Dict:
        """转换图片格式到标准格式"""
        if not PIL_AVAILABLE:
            return img_info
        
        try:
            import io
            img = Image.open(io.BytesIO(img_info['data']))
            
            # 转换为目标格式
            output = io.BytesIO()
            if target_format == '.png':
                img.save(output, format='PNG')
            elif target_format in ['.jpg', '.jpeg']:
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')
                img.save(output, format='JPEG', quality=95)
            
            img_info['data'] = output.getvalue()
            img_info['format'] = target_format
            
        except Exception as e:
            self.log(f"格式转换失败: {str(e)}", "WARNING")
        
        return img_info
    
    def _get_image_metadata(self, filepath: Path, image_data: bytes) -> Dict:
        """获取图片元数据"""
        metadata = {
            'width': 0,
            'height': 0,
            'dpi': 0
        }
        
        if PIL_AVAILABLE:
            try:
                import io
                img = Image.open(io.BytesIO(image_data))
                metadata['width'] = img.width
                metadata['height'] = img.height
                
                # 获取 DPI
                if 'dpi' in img.info:
                    metadata['dpi'] = img.info['dpi'][0]
                    
            except Exception:
                pass
        
        return metadata
    
    def generate_index(self, material_name: str, images: List[Dict], 
                      paragraphs: Dict, docx_path: str) -> Dict:
        """
        生成素材配图索引
        
        Args:
            material_name: 素材名称
            images: 提取的图片列表
            paragraphs: 段落关联信息
            docx_path: 原始 docx 路径
            
        Returns:
            索引字典
        """
        index = {
            'material_name': material_name,
            'original_file': str(docx_path),
            'extraction_time': datetime.now().isoformat(),
            'total_images': len(images),
            'image_types': {},
            'images': images,
            'paragraphs': paragraphs
        }
        
        # 统计图片类型
        for img in images:
            img_type = img.get('type', '其他')
            index['image_types'][img_type] = index['image_types'].get(img_type, 0) + 1
        
        return index
    
    def save_index(self, index: Dict, output_file: str = None):
        """保存索引到 Markdown 文件"""
        if output_file is None:
            output_file = self.output_dir.parent / 'iteration' / 'material-index.md'
        else:
            output_file = Path(output_file)
        
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # 生成 Markdown 内容
        md_content = self._generate_markdown_index(index)
        
        # 追加或创建文件
        if output_file.exists():
            with open(output_file, 'a', encoding='utf-8') as f:
                f.write('\n' + md_content)
        else:
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write("# 素材配图索引\n\n")
                f.write("本索引记录所有从素材文档中提取的配图信息。\n\n")
                f.write(md_content)
        
        self.log(f"索引已保存: {output_file}")
    
    def _generate_markdown_index(self, index: Dict) -> str:
        """生成 Markdown 格式的索引内容"""
        material_name = index['material_name']
        
        md = f"\n## 素材: {material_name}\n\n"
        md += f"**原始文件**: {index['original_file']}\n\n"
        md += f"**提取时间**: {index['extraction_time']}\n\n"
        md += f"**配图总数**: {index['total_images']}\n\n"
        
        # 类型统计
        if index['image_types']:
            md += "### 配图类型分布\n\n"
            md += "| 类型 | 数量 |\n"
            md += "|------|------|\n"
            for img_type, count in index['image_types'].items():
                md += f"| {img_type} | {count} |\n"
            md += "\n"
        
        # 配图明细
        md += "### 配图明细\n\n"
        md += "| 图片ID | 文件名 | 类型 | 尺寸 | 大小 |\n"
        md += "|--------|--------|------|------|------|\n"
        
        for img in index['images']:
            size_kb = img.get('size_bytes', 0) / 1024
            dimensions = f"{img.get('width', 0)}x{img.get('height', 0)}"
            md += f"| {img['id']} | {img['filename']} | {img['type']} | {dimensions} | {size_kb:.1f}KB |\n"
        
        md += "\n"
        
        return md
    
    def process_material(self, docx_path: str) -> Dict:
        """
        处理单个素材文件
        
        Args:
            docx_path: docx 文件路径
            
        Returns:
            处理结果字典
        """
        docx_path = Path(docx_path)
        material_name = docx_path.stem
        
        self.log(f"处理素材: {material_name}")
        
        # 提取图片
        images = self.extract_images_from_docx(str(docx_path))
        
        if not images:
            self.log(f"未找到图片: {material_name}", "WARNING")
            return {
                'material_name': material_name,
                'success': False,
                'images': [],
                'error': 'No images found'
            }
        
        # 识别段落关联
        paragraphs = self.identify_paragraph_relationship(str(docx_path))
        
        # 保存图片
        saved_images = self.save_images(images, material_name)
        
        # 生成索引
        index = self.generate_index(material_name, saved_images, paragraphs, docx_path)
        
        # 保存索引
        self.save_index(index)
        
        self.log(f"素材处理完成: {material_name} ({len(saved_images)} 张图片)")
        
        return {
            'material_name': material_name,
            'success': True,
            'images': saved_images,
            'index': index
        }
    
    def process_directory(self, material_dir: str) -> List[Dict]:
        """
        批量处理目录中的所有 docx 文件
        
        Args:
            material_dir: 素材目录路径
            
        Returns:
            处理结果列表
        """
        material_dir = Path(material_dir)
        results = []
        
        if not material_dir.exists():
            self.log(f"目录不存在: {material_dir}", "ERROR")
            return results
        
        # 查找所有 docx 文件
        docx_files = list(material_dir.glob('**/*.docx'))
        
        self.log(f"找到 {len(docx_files)} 个 docx 文件")
        
        for docx_file in docx_files:
            result = self.process_material(str(docx_file))
            results.append(result)
        
        # 生成汇总报告
        self._generate_summary_report(results)
        
        return results
    
    def _generate_summary_report(self, results: List[Dict]):
        """生成处理汇总报告"""
        total_materials = len(results)
        successful = sum(1 for r in results if r['success'])
        total_images = sum(len(r['images']) for r in results)
        
        self.log("=" * 60)
        self.log("处理汇总")
        self.log("=" * 60)
        self.log(f"素材文件总数: {total_materials}")
        self.log(f"成功处理: {successful}")
        self.log(f"失败: {total_materials - successful}")
        self.log(f"提取配图总数: {total_images}")
        self.log("=" * 60)


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='从 docx 素材文件中提取配图',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  %(prog)s --material 生态修复技术导则.docx
  %(prog)s --material-dir ./materials/ --verbose
  %(prog)s --material F:/docs/导则.docx --output assets/docx_images/
        '''
    )
    
    parser.add_argument('--material', '-m',
                      help='单个 docx 素材文件路径')
    parser.add_argument('--material-dir', '-d',
                      help='素材目录路径（批量处理）')
    parser.add_argument('--output', '-o',
                      default='assets/docx_images',
                      help='输出目录 (默认: assets/docx_images)')
    parser.add_argument('--index', '-i',
                      default='iteration/material-index.md',
                      help='索引文件路径 (默认: iteration/material-index.md)')
    parser.add_argument('--verbose', '-v',
                      action='store_true',
                      help='显示详细日志')
    
    args = parser.parse_args()
    
    # 检查参数
    if not args.material and not args.material_dir:
        parser.print_help()
        sys.exit(1)
    
    # 创建提取器
    extractor = MaterialImageExtractor(
        output_dir=args.output,
        verbose=args.verbose
    )
    
    # 处理素材
    if args.material:
        result = extractor.process_material(args.material)
        if result['success']:
            print(f"\n✓ 成功提取 {len(result['images'])} 张图片")
            sys.exit(0)
        else:
            print(f"\n✗ 提取失败: {result.get('error', 'Unknown error')}")
            sys.exit(1)
    
    elif args.material_dir:
        results = extractor.process_directory(args.material_dir)
        successful = sum(1 for r in results if r['success'])
        print(f"\n✓ 成功处理 {successful}/{len(results)} 个素材文件")
        sys.exit(0 if successful > 0 else 1)


if __name__ == '__main__':
    main()
