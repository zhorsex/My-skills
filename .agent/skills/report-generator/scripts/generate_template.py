#!/usr/bin/env python3
"""
Template Generation Script for Report Generator

Generates Word document templates with customizable parameters.
Supports predefined templates and custom template generation.

Usage:
    python generate_template.py --type geological --output templates/geo.docx
    python generate_template.py --type custom --interactive
    python generate_template.py --preset "my-format"
"""

import argparse
import json
import os
import sys
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# Try to import docx
try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Error: python-docx library not found. Install with: pip install python-docx")
    sys.exit(1)


# Default configurations
DEFAULT_CONFIG = {
    "page": {
        "paper_size": "A4",
        "margin_top": 2.75,
        "margin_bottom": 2.75,
        "margin_left": 2.86,  # Inner (binding side)
        "margin_right": 2.27,  # Outer
        "gutter": 1.0,
        "header_distance": 1.5,
        "footer_distance": 1.75
    },
    "fonts": {
        "chinese": "宋体",
        "english": "Times New Roman",
        "cover_title": "黑体",
        "chapter_title": "宋体",
        "section_title": "黑体",
        "subsection_title": "宋体",
        "body": "宋体"
    },
    "sizes": {
        "cover_subtitle": 15,  #小三号
        "cover_title": 26,     #小一号
        "cover_unit": 16,      #三号
        "cover_date": 15,      #小三号
        "chapter": 22,         #二号
        "section": 16,         #三号
        "subsection": 14,    #四号
        "body": 10.5,          #五号
        "table_header": 10.5,
        "table_content": 9,    #小五号
        "table_note": 7.5,     #六号
        "figure": 9,           #小五号
        "reference": 9         #小五号
    },
    "paragraph": {
        "line_spacing": 1.25,
        "first_line_indent": 0.74,  # 2 characters in cm
        "chapter_before": 4,
        "chapter_after": 4,
        "section_before": 0.5,
        "section_after": 0.25
    },
    "elements": {
        "cover": True,
        "title_page": True,
        "toc": True,
        "header_footer": True,
        "page_numbers": True,
        "tables": True,
        "figures": True,
        "references": True
    }
}

# Predefined template configurations
PREDEFINED_TEMPLATES = {
    "geological": {
        "name": "地质资料归档格式模板",
        "description": "符合河南省国土资源厅豫国土资发［2006］144号文标准",
        "config": DEFAULT_CONFIG
    },
    "engineering": {
        "name": "通用工程技术报告模板",
        "description": "简洁通用的工程报告格式",
        "config": {
            **DEFAULT_CONFIG,
            "fonts": {
                **DEFAULT_CONFIG["fonts"],
                "chinese": "宋体",
                "cover_title": "黑体",
                "chapter_title": "黑体"
            },
            "sizes": {
                **DEFAULT_CONFIG["sizes"],
                "chapter": 18,
                "section": 14,
                "body": 12
            }
        }
    }
}


class TemplateGenerator:
    """Generates Word document templates with customizable parameters."""
    
    def __init__(self, config: Dict):
        self.config = config
        self.doc = None
        
    def create_document(self) -> Document:
        """Create a new Word document with template structure."""
        self.doc = Document()
        self._setup_page()
        self._create_styles()
        self._add_sample_content()
        return self.doc
    
    def _setup_page(self):
        """Configure page settings."""
        section = self.doc.sections[0]
        page = self.config["page"]
        
        # Set margins (convert cm to inches for python-docx)
        section.top_margin = Cm(page["margin_top"])
        section.bottom_margin = Cm(page["margin_bottom"])
        section.left_margin = Cm(page["margin_left"])
        section.right_margin = Cm(page["margin_right"])
        section.gutter = Cm(page["gutter"])
        section.header_distance = Cm(page["header_distance"])
        section.footer_distance = Cm(page["footer_distance"])
        
        # Set different first page (for cover)
        section.different_first_page_header_footer = True
        
    def _create_styles(self):
        """Create all paragraph styles for the template."""
        fonts = self.config["fonts"]
        sizes = self.config["sizes"]
        para = self.config["paragraph"]
        
        # Cover styles
        self._create_style("CoverSubtitle", fonts["chinese"], sizes["cover_subtitle"], 
                          bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        self._create_style("CoverTitle", fonts["cover_title"], sizes["cover_title"],
                          bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._create_style("CoverUnit", fonts["chinese"], sizes["cover_unit"],
                          bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._create_style("CoverDate", fonts["chinese"], sizes["cover_date"],
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Title page styles
        self._create_style("TitlePageTitle", fonts["chinese"], sizes["chapter"],
                          bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._create_style("TitlePageInfo", fonts["chinese"], sizes["section"],
                          alignment=WD_ALIGN_PARAGRAPH.LEFT, left_indent=3.7)
        self._create_style("TitlePagePersonnel", fonts["chinese"], sizes["section"],
                          alignment=WD_ALIGN_PARAGRAPH.LEFT, left_indent=2.5)
        
        # TOC styles
        self._create_style("TocTitle", fonts["chinese"], sizes["section"],
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._create_style("TocChapter", fonts["chinese"], sizes["body"],
                          bold=True)
        self._create_style("TocSection", fonts["chinese"], sizes["body"],
                          left_indent=0.74)
        
        # Content styles
        self._create_style("ChapterTitle", fonts["chapter_title"], sizes["chapter"],
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          space_before=para["chapter_before"],
                          space_after=para["chapter_after"])
        self._create_style("SectionTitle", fonts["section_title"], sizes["section"],
                          bold=True, space_before=para["section_before"],
                          space_after=para["section_after"])
        self._create_style("SubsectionTitle", fonts["subsection_title"], sizes["subsection"],
                          bold=True)
        self._create_style("BodyText", fonts["body"], sizes["body"],
                          first_line_indent=para["first_line_indent"],
                          line_spacing=para["line_spacing"])
        
        # Table styles
        self._create_style("TableNumber", fonts["english"], sizes["table_header"],
                          bold=True, left_indent=0.74)
        self._create_style("TableTitle", fonts["chinese"], sizes["table_header"],
                          bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._create_style("TableContent", fonts["chinese"], sizes["table_content"],
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._create_style("TableNote", fonts["chinese"], sizes["table_note"],
                          left_indent=0.37)
        
        # Figure styles
        self._create_style("FigureTitle", fonts["chinese"], sizes["figure"],
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Reference styles
        self._create_style("ReferenceTitle", fonts["chinese"], sizes["body"],
                          bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self._create_style("ReferenceContent", fonts["chinese"], sizes["reference"])
        
    def _create_style(self, name: str, font_cn: str, size_pt: float, 
                     bold: bool = False, alignment=None,
                     left_indent: float = 0, first_line_indent: float = 0,
                     space_before: float = 0, space_after: float = 0,
                     line_spacing: float = None):
        """Create a paragraph style."""
        style = self.doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        
        # Font settings
        font = style.font
        font.name = "Times New Roman"  # For English/numbers
        font.size = Pt(size_pt)
        font.bold = bold
        
        # Chinese font
        run_font = style.element.rPr
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), font_cn)
        run_font.append(rFonts)
        
        # Paragraph settings
        para_format = style.paragraph_format
        if alignment:
            para_format.alignment = alignment
        if left_indent:
            para_format.left_indent = Cm(left_indent)
        if first_line_indent:
            para_format.first_line_indent = Cm(first_line_indent)
        if space_before:
            para_format.space_before = Pt(space_before * 12)  # Approximate
        if space_after:
            para_format.space_after = Pt(space_after * 12)
        if line_spacing:
            para_format.line_spacing = line_spacing
            
    def _add_sample_content(self):
        """Add sample content demonstrating all styles."""
        elements = self.config["elements"]
        
        # Cover page
        if elements["cover"]:
            self._add_cover_page()
            self.doc.add_page_break()
        
        # Title page
        if elements["title_page"]:
            self._add_title_page()
            self.doc.add_page_break()
        
        # TOC
        if elements["toc"]:
            self._add_toc()
            self.doc.add_page_break()
        
        # Sample chapters
        self._add_sample_chapters()
        
    def _add_cover_page(self):
        """Add cover page with sample content."""
        # Subtitle
        p = self.doc.add_paragraph("地质调查项目成果报告", style="CoverSubtitle")
        
        # Empty lines
        for _ in range(4):
            self.doc.add_paragraph()
        
        # Title
        p = self.doc.add_paragraph("成果报告题名", style="CoverTitle")
        
        # Empty lines
        for _ in range(5):
            self.doc.add_paragraph()
        
        # Unit
        p = self.doc.add_paragraph("项目承担单位", style="CoverUnit")
        
        # Empty lines
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Date
        p = self.doc.add_paragraph("二〇二六年一月", style="CoverDate")
        
    def _add_title_page(self):
        """Add title page with project information."""
        # Title
        p = self.doc.add_paragraph("成果报告题名", style="TitlePageTitle")
        
        # Empty lines
        for _ in range(2):
            self.doc.add_paragraph()
        
        # Project info
        p = self.doc.add_paragraph("项目编码：XXXXXXXX", style="TitlePageInfo")
        p = self.doc.add_paragraph("任务书编号：XXXXXXXX", style="TitlePageInfo")
        p = self.doc.add_paragraph("工作起止时间：2024.01-2025.12", style="TitlePageInfo")
        
        # Empty lines
        for _ in range(2):
            self.doc.add_paragraph()
        
        # Personnel
        p = self.doc.add_paragraph("项目负责人：XXX", style="TitlePagePersonnel")
        p = self.doc.add_paragraph("报告主编：XXX", style="TitlePagePersonnel")
        p = self.doc.add_paragraph("技术负责：XXX", style="TitlePagePersonnel")
        p = self.doc.add_paragraph("主要编写人：XXX、XXX、XXX", style="TitlePagePersonnel")
        p = self.doc.add_paragraph("实施单位：XXX地质调查院", style="TitlePagePersonnel")
        
    def _add_toc(self):
        """Add table of contents."""
        # TOC Title
        p = self.doc.add_paragraph("目  录", style="TocTitle")
        
        # TOC Entries
        p = self.doc.add_paragraph("1 绪言", style="TocChapter")
        p = self.doc.add_paragraph("1.1 目的任务", style="TocSection")
        p = self.doc.add_paragraph("1.2 位置交通", style="TocSection")
        
        p = self.doc.add_paragraph("2 区域地质背景", style="TocChapter")
        p = self.doc.add_paragraph("2.1 地层", style="TocSection")
        
        p = self.doc.add_paragraph("参考文献", style="TocChapter")
        
    def _add_sample_chapters(self):
        """Add sample chapter content."""
        # Chapter 1
        p = self.doc.add_paragraph("1 绪言", style="ChapterTitle")
        
        p = self.doc.add_paragraph("1.1 目的任务", style="SectionTitle")
        
        p = self.doc.add_paragraph(
            "本章节为示例内容，展示正文格式。根据《成果地质资料电子文件汇交格式要求》，"
            "正文内容使用五号宋体，首行缩进2字符，行距为多倍行距1.25。",
            style="BodyText"
        )
        
        p = self.doc.add_paragraph(
            "数字与单位之间需要加空格，如：10 m、1000 km²。范围符号使用示例：10～20 m，"
            "5%～10%，600～800 Ma。中文地名范围使用一字线：郑州—开封—商丘。",
            style="BodyText"
        )
        
        # Sample table
        if self.config["elements"]["tables"]:
            p = self.doc.add_paragraph("表1-1", style="TableNumber")
            p = self.doc.add_paragraph("示例数据表", style="TableTitle")
            
            table = self.doc.add_table(rows=3, cols=3)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "参数"
            hdr_cells[1].text = "数值"
            hdr_cells[2].text = "单位"
            
            # Data rows
            row1_cells = table.rows[1].cells
            row1_cells[0].text = "长度"
            row1_cells[1].text = "100"
            row1_cells[2].text = "m"
            
            row2_cells = table.rows[2].cells
            row2_cells[0].text = "面积"
            row2_cells[1].text = "50"
            row2_cells[2].text = "km²"
            
            p = self.doc.add_paragraph("注：表注使用六号宋体", style="TableNote")
        
        # Chapter 2
        p = self.doc.add_paragraph("2 区域地质背景", style="ChapterTitle")
        p = self.doc.add_paragraph("2.1 地层", style="SectionTitle")
        
        p = self.doc.add_paragraph(
            "第二章示例内容。章与章之间使用分页功能。比例尺符号使用∶，不要用冒号：代替。",
            style="BodyText"
        )
        
        # References
        if self.config["elements"]["references"]:
            p = self.doc.add_paragraph("参考文献", style="ReferenceTitle")
            p = self.doc.add_paragraph(
                "[1] 河南省国土资源厅. 成果地质资料电子文件汇交格式要求[S]. 2006.",
                style="ReferenceContent"
            )


def interactive_mode():
    """Run interactive mode for custom template generation."""
    print("\n" + "="*60)
    print("       自定义Word模板生成")
    print("="*60 + "\n")
    
    config = DEFAULT_CONFIG.copy()
    
    # Page settings
    print("【页面设置】")
    print(f"纸张大小 [A4]: ", end="")
    input_val = input().strip()
    
    print(f"上页边距 (cm) [{config['page']['margin_top']}]: ", end="")
    val = input().strip()
    if val: config['page']['margin_top'] = float(val)
    
    print(f"下页边距 (cm) [{config['page']['margin_bottom']}]: ", end="")
    val = input().strip()
    if val: config['page']['margin_bottom'] = float(val)
    
    print(f"左页边距（内侧）(cm) [{config['page']['margin_left']}]: ", end="")
    val = input().strip()
    if val: config['page']['margin_left'] = float(val)
    
    print(f"右页边距（外侧）(cm) [{config['page']['margin_right']}]: ", end="")
    val = input().strip()
    if val: config['page']['margin_right'] = float(val)
    
    print(f"装订线 (cm) [{config['page']['gutter']}]: ", end="")
    val = input().strip()
    if val: config['page']['gutter'] = float(val)
    
    # Font settings
    print("\n【字体设置】")
    print(f"中文字体 [{config['fonts']['chinese']}]: ", end="")
    val = input().strip()
    if val: config['fonts']['chinese'] = val
    
    print(f"章标题字体 [{config['fonts']['chapter_title']}]: ", end="")
    val = input().strip()
    if val: config['fonts']['chapter_title'] = val
    
    # Size settings
    print("\n【字号设置】")
    print(f"章标题字号 (pt) [{config['sizes']['chapter']}]: ", end="")
    val = input().strip()
    if val: config['sizes']['chapter'] = float(val)
    
    print(f"正文字号 (pt) [{config['sizes']['body']}]: ", end="")
    val = input().strip()
    if val: config['sizes']['body'] = float(val)
    
    # Output
    print("\n【输出设置】")
    print("输出文件名 [custom-template.docx]: ", end="")
    filename = input().strip()
    if not filename:
        filename = "custom-template.docx"
    
    # Preview
    print("\n" + "="*60)
    print("参数预览：")
    print(f"  页边距: 上{config['page']['margin_top']}/下{config['page']['margin_bottom']}/"
          f"左{config['page']['margin_left']}/右{config['page']['margin_right']} cm")
    print(f"  字体: {config['fonts']['chinese']}")
    print(f"  章标题: {config['fonts']['chapter_title']} {config['sizes']['chapter']}pt")
    print(f"  正文: {config['fonts']['body']} {config['sizes']['body']}pt")
    print("="*60)
    
    print("\n确认生成模板? (Y/n): ", end="")
    confirm = input().strip().lower()
    
    if confirm in ('', 'y', 'yes'):
        return config, filename
    else:
        print("已取消")
        return None, None


def save_preset(name: str, config: Dict):
    """Save configuration as a named preset."""
    presets_dir = Path("templates/.presets")
    presets_dir.mkdir(parents=True, exist_ok=True)
    
    preset_file = presets_dir / f"{name}.json"
    with open(preset_file, 'w', encoding='utf-8') as f:
        json.dump(config, f, ensure_ascii=False, indent=2)
    
    print(f"预设已保存: {preset_file}")


def load_preset(name: str) -> Optional[Dict]:
    """Load a named preset."""
    preset_file = Path(f"templates/.presets/{name}.json")
    
    if not preset_file.exists():
        return None
    
    with open(preset_file, 'r', encoding='utf-8') as f:
        return json.load(f)


def list_presets():
    """List all available presets."""
    presets_dir = Path("templates/.presets")
    
    if not presets_dir.exists():
        print("没有保存的预设")
        return
    
    presets = list(presets_dir.glob("*.json"))
    
    if not presets:
        print("没有保存的预设")
        return
    
    print("\n可用预设：")
    for i, preset in enumerate(presets, 1):
        print(f"  {i}. {preset.stem}")


def main():
    parser = argparse.ArgumentParser(
        description="生成Word文档模板",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 生成地质资料归档格式模板
  python generate_template.py --type geological --output templates/geo.docx
  
  # 交互式自定义生成
  python generate_template.py --type custom --interactive
  
  # 使用预设
  python generate_template.py --preset my-format --output templates/my.docx
        """
    )
    
    parser.add_argument('--type', choices=['geological', 'engineering', 'custom'],
                       help='模板类型')
    parser.add_argument('--output', '-o', help='输出文件路径')
    parser.add_argument('--interactive', '-i', action='store_true',
                       help='交互式模式（用于custom类型）')
    parser.add_argument('--preset', help='使用保存的预设')
    parser.add_argument('--list-presets', action='store_true',
                       help='列出所有预设')
    
    # Custom parameters
    parser.add_argument('--page-margin', help='页边距 (上,下,左,右)，如：2.75,2.75,2.86,2.27')
    parser.add_argument('--gutter', type=float, help='装订线 (cm)')
    parser.add_argument('--body-font', help='正文字体')
    parser.add_argument('--body-size', type=float, help='正文字号 (pt)')
    
    args = parser.parse_args()
    
    # List presets
    if args.list_presets:
        list_presets()
        return
    
    # Load preset
    if args.preset:
        config = load_preset(args.preset)
        if config is None:
            print(f"错误: 找不到预设 '{args.preset}'")
            print("使用 --list-presets 查看可用预设")
            return
        output = args.output or f"templates/{args.preset}.docx"
    
    # Predefined templates
    elif args.type in PREDEFINED_TEMPLATES:
        template_info = PREDEFINED_TEMPLATES[args.type]
        config = template_info["config"]
        print(f"生成模板: {template_info['name']}")
        print(f"说明: {template_info['description']}")
        
        default_name = f"{args.type}-template.docx"
        output = args.output or f"templates/{default_name}"
    
    # Custom template
    elif args.type == 'custom':
        if args.interactive:
            config, filename = interactive_mode()
            if config is None:
                return
            output = args.output or f"templates/{filename}"
        else:
            # Use command line parameters
            config = DEFAULT_CONFIG.copy()
            
            if args.page_margin:
                margins = [float(x) for x in args.page_margin.split(',')]
                config['page']['margin_top'] = margins[0]
                config['page']['margin_bottom'] = margins[1]
                config['page']['margin_left'] = margins[2]
                config['page']['margin_right'] = margins[3]
            
            if args.gutter:
                config['page']['gutter'] = args.gutter
            
            if args.body_font:
                config['fonts']['body'] = args.body_font
                config['fonts']['chinese'] = args.body_font
            
            if args.body_size:
                config['sizes']['body'] = args.body_size
            
            output = args.output or "templates/custom-template.docx"
    
    else:
        parser.print_help()
        return
    
    # Generate template
    print(f"\n正在生成模板...")
    print(f"输出文件: {output}")
    
    try:
        generator = TemplateGenerator(config)
        doc = generator.create_document()
        
        # Ensure output directory exists
        output_path = Path(output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        doc.save(output)
        print(f"✓ 模板生成成功: {output}")
        
        # Ask to save as preset (for custom templates)
        if args.type == 'custom' and args.interactive:
            print("\n是否保存为预设供将来使用? (y/N): ", end="")
            save = input().strip().lower()
            if save in ('y', 'yes'):
                print("预设名称: ", end="")
                preset_name = input().strip()
                if preset_name:
                    save_preset(preset_name, config)
    
    except Exception as e:
        print(f"错误: 生成模板失败 - {e}")
        import traceback
        traceback.print_exc()
        return 1
    
    return 0


if __name__ == '__main__':
    sys.exit(main())
